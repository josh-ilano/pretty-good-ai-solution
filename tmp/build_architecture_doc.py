from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Pretty_Good_AI_Architecture.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "17202A"
MUTED = "5C6773"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "EAF2F8"
LIGHT_GOLD = "FFF7E0"
WHITE = "FFFFFF"
GREEN = "247A4D"
AMBER = "8A5A00"


def set_run_font(run, size=None, bold=None, italic=None, color=INK):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_table_borders(table, color="D0D5DD", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        remainder = paragraph.add_run(text[len(bold_lead):])
        set_run_font(remainder)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_bullet(doc, text, level=0):
    paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), BLUE)
    left.set(qn("w:space"), "8")
    borders.append(left)
    p_pr.append(borders)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, bold=True, color=DARK_BLUE)
    text_run = paragraph.add_run(text)
    set_run_font(text_run)
    return paragraph


def set_cell_text(cell, text, bold=False, color=INK, size=9.2, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_matrix(doc, headers, rows, widths, status_column=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for index, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[index], LIGHT_GRAY)
        set_cell_text(table.rows[0].cells[index], header, bold=True, color=DARK_BLUE, size=9.3)
    for row_values in rows:
        row = table.add_row()
        set_row_cant_split(row)
        for index, value in enumerate(row_values):
            is_status = status_column and index == len(row_values) - 1
            alignment = WD_ALIGN_PARAGRAPH.CENTER if is_status else WD_ALIGN_PARAGRAPH.LEFT
            color = INK
            if is_status:
                color = GREEN if value in {"Implemented", "Ready"} else AMBER
            set_cell_text(row.cells[index], value, bold=is_status, color=color, alignment=alignment)
    set_table_geometry(table, widths)
    return table


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    left = header.add_run("Pretty Good AI Voicebot Evaluator")
    set_run_font(left, size=9, bold=True, color=MUTED)
    right = header.add_run("\tArchitecture")
    set_run_font(right, size=9, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    doc.core_properties.title = "Pretty Good AI Voicebot Evaluator - Architecture"
    doc.core_properties.subject = "AI Engineering Challenge architecture and requirements traceability"
    doc.core_properties.author = "Project Team"
    doc.core_properties.keywords = "voicebot, SignalWire, OpenAI Realtime, RAG, architecture"

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(5)
    kicker_run = kicker.add_run("TECHNICAL ARCHITECTURE BRIEF")
    set_run_font(kicker_run, size=9.5, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("Pretty Good AI Voicebot Evaluator")
    set_run_font(title_run, size=24, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run("Architecture, design decisions, and challenge requirements traceability")
    set_run_font(subtitle_run, size=13, color=MUTED)

    metadata = [
        ("Purpose", "AI Engineering Challenge submission"),
        ("Scope", "Automated fictional-patient calls to the authorized assessment line"),
        ("Version", "1.0 - August 20, 2026"),
        ("Status", "Editable submission draft"),
        ("Source", "Pretty Good AI - AI Engineering Challenge (4 pages) and repository review"),
    ]
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}: ")
        set_run_font(label_run, size=10.5, bold=True, color=DARK_BLUE)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=10.5)

    add_callout(
        doc,
        "Scope boundary",
        "All automated calls are restricted in code to +1-805-439-8008. Scenarios use fictional data and are intended only for the authorized assessment.",
    )

    add_heading(doc, "1. Architecture overview", 1)
    add_body(
        doc,
        "The system separates prompt generation, live communication, and artifact capture into three independently testable layers. The prompt layer uses a local SQLite FTS5 policy index to retrieve page-aware evidence and build a bounded fictional-patient scenario, or it accepts a fully manual prompt. The communication layer in realtime_call.py creates exactly one guarded SignalWire call to the assessment number and opens a bidirectional PCMU media stream to the OpenAI Realtime API. Server-side voice activity detection identifies office turns, while application-level menu filtering, turn settling, response buffering, and playback marks keep the conversation coherent and prevent partial or prerecorded audio from producing premature replies.",
    )
    add_body(
        doc,
        "The capture layer in capture_call.py owns the call-status and recording-status webhooks, timestamped transcripts, authenticated recording download, retries, and per-run artifact directories. This boundary keeps telephony orchestration small while ensuring that each generated scenario, transcript, and MP3 can be traced to a single test case. The design favors a direct Realtime audio bridge over batch speech pipelines for natural turn-taking; local lexical retrieval over hosted vector infrastructure for deterministic setup; and a JSON scenario contract over prompt-only generation so policy evidence, expected safe behavior, and observable failure conditions remain auditable after the call.",
    )

    doc.add_page_break()
    add_heading(doc, "2. Challenge requirements traceability", 1)
    add_body(doc, "The matrix distinguishes implemented capabilities from submission activities that still require human completion.")
    requirements = [
        ("Authorized calling", "Call only +1-805-439-8008", "Hard-coded destination plus scenario-contract validation; source number is configured separately.", "Implemented"),
        ("Patient simulation", "Realistic scenarios across scheduling, refills, questions, and edge cases", "RAG generator provides 31 situations across 30 categories; manual prompts are also supported.", "Implemented"),
        ("Recording", "Capture both sides in MP3 or OGG", "SignalWire dual-channel recording is downloaded as recording.mp3 after finalization.", "Implemented"),
        ("Transcription", "Capture both sides of each conversation", "Office and played patient turns are timestamped in transcript.txt.", "Implemented"),
        ("Bug discovery", "Identify quality or policy failures", "Scenario JSON stores evidence, expected safe behavior, and failure conditions for transcript review.", "Implemented"),
        ("Run instructions", "Single command after setup", "README documents python realtime_call.py --rag after credentials and tunnel setup.", "Ready"),
        ("Architecture", "Explain operation and design choices in 1-2 paragraphs", "Section 1 contains the required two-paragraph explanation; later sections provide traceability.", "Ready"),
        ("Call evidence", "At least 10 complete, 1-3 minute calls", "Runtime creates paired artifacts; submission owner must confirm ten complete pairs and call quality.", "Complete before submission"),
        ("Bug report", "Document useful issues with call references", "Use JSON criteria and transcript timestamps; final consolidated report remains a submission task.", "Complete before submission"),
        ("Loom videos", "Walkthrough plus AI-debugging recording", "External recordings must be created, made public, and checked against the time/content requirements.", "Complete before submission"),
        ("Repository", "Public GitHub repository with no secrets", "Code, README, requirements, and .env.example are present; public visibility and secret scan require final confirmation.", "Verify before submission"),
        ("Originating number", "Use one source number and submit it in E.164", "SIGNALWIRE_FROM_NUMBER supplies one source number for all runs; record it accurately in the form.", "Verify before submission"),
    ]
    add_heading(doc, "2.1 Implemented system capabilities", 2)
    add_matrix(
        doc,
        ("Area", "Challenge requirement", "Implementation / evidence", "Status"),
        requirements[:7],
        (1450, 2250, 4060, 1600),
        status_column=True,
    )

    doc.add_page_break()
    add_heading(doc, "2.2 Submission deliverables and external checks", 2)
    add_matrix(
        doc,
        ("Area", "Challenge requirement", "Implementation / evidence", "Status"),
        requirements[7:],
        (1450, 2250, 4060, 1600),
        status_column=True,
    )

    doc.add_page_break()
    add_heading(doc, "3. Component responsibilities", 1)
    components = [
        ("realtime_call.py", "Communication and control", "Loads a prompt, validates configuration, starts one guarded call, serves cXML/websocket endpoints, and bridges audio."),
        ("scenario_generator/", "Prompt and evaluation layer", "Selects a diverse policy conflict, retrieves evidence, creates auditable JSON, and validates manual prompt contracts."),
        ("rag_pipeline/", "Policy knowledge layer", "Extracts page-aware PDF chunks and exposes BM25 retrieval through a local SQLite FTS5 index."),
        ("capture_call.py", "Evidence capture", "Persists played patient turns and office transcripts, receives provider webhooks, and downloads the finalized MP3."),
        ("SignalWire", "Telephony provider", "Places the PSTN call, streams PCMU audio, emits call/recording events, and hosts recording media."),
        ("OpenAI Realtime API", "Conversational patient", "Consumes office audio and produces natural spoken fictional-patient responses under the selected prompt."),
    ]
    add_matrix(doc, ("Component", "Role", "Responsibility"), components, (2100, 2200, 5060))

    add_heading(doc, "4. Runtime and data flow", 1)
    flows = [
        "At startup, the operator chooses RAG mode or supplies a complete manual patient prompt.",
        "RAG mode excludes recently used topics, retrieves relevant policy chunks, and saves a uniquely identified JSON scenario under generated_prompt/.",
        "The scenario contract validates the fictional-data marker and authorized destination before its patient_prompt enters the communication layer.",
        "SignalWire places one outbound call and requests /cxml, which instructs it to open a bidirectional PCMU stream to /media-stream.",
        "The application forwards office audio to OpenAI Realtime, filters recorded menus and incomplete turns, and returns complete patient utterances to SignalWire.",
        "Playback acknowledgements determine which generated patient turns are committed to the transcript; office turns are transcribed from input audio.",
        "SignalWire sends completion webhooks. The capture layer returns HTTP 204 immediately, downloads the MP3 asynchronously with retries, and reports both artifacts ready.",
    ]
    for item in flows:
        add_number(doc, item)

    doc.add_page_break()
    add_heading(doc, "5. Key design decisions and tradeoffs", 1)
    decisions = [
        ("Realtime audio bridge", "Produces natural latency and turn-taking, which the challenge evaluates before code quality.", "Requires websocket state, audio buffering, VAD tuning, and careful interruption handling."),
        ("SignalWire compatibility API", "Combines outbound calling, bidirectional streaming, status callbacks, and recordings.", "Requires a public HTTPS tunnel and provider credentials during local development."),
        ("Local SQLite FTS5 retrieval", "Keeps ingestion reproducible, fast, inexpensive, and free of hosted vector-database setup.", "Lexical BM25 search has weaker semantic matching than an embedding-based retriever."),
        ("Layered JSON contract", "Preserves prompt, evidence, expected behavior, and failure criteria for repeatable analysis.", "Creates more artifacts and schema maintenance than directly interpolating a prompt."),
        ("Playback-confirmed transcript", "Records only patient speech acknowledged as played to the remote agent.", "Adds mark tracking and buffering complexity but avoids logging unheard generated text."),
        ("Hard-coded destination guard", "Prevents accidental calls outside the assessment scope.", "Intentionally limits reuse as a general-purpose calling platform."),
    ]
    add_matrix(doc, ("Decision", "Why selected", "Tradeoff"), decisions, (2200, 3550, 3610))

    doc.add_page_break()
    add_heading(doc, "6. Reliability, safety, and evidence quality", 1)
    add_bullet(doc, "The destination is fixed to the authorized assessment number, and the configured SignalWire source number cannot equal that destination.")
    add_bullet(doc, "API keys remain in .env; .env.example documents required variables without credentials.")
    add_bullet(doc, "Generated scenarios require fictional data and impose bounded attempts, final-outcome stopping rules, and no fabricated emergencies.")
    add_bullet(doc, "Recorded menus and short transcript fragments do not trigger the patient; a settling window reduces false end-of-turn responses.")
    add_bullet(doc, "Dual-channel recording, playback marks, timestamps, provider status monitoring, and download retries improve evidence completeness.")
    add_bullet(doc, "Ctrl+T provides an operator-controlled early stop while preserving time for recording finalization.")
    add_callout(
        doc,
        "Not a production healthcare system",
        "The architecture is an assessment harness. It is not designed or represented as a HIPAA production deployment, a clinical decision system, or a tool for real patient data.",
        fill=LIGHT_GOLD,
    )

    add_heading(doc, "7. Verification and submission readiness", 1)
    add_body(doc, "Automated checks validate scenario classification, diversity history, destination enforcement, prompt contracts, ingestion, chunking, and retrieval. These tests do not place phone calls. Runtime evidence still requires listening to each recording and reviewing its transcript because conversational quality is the challenge's first evaluation gate.")
    add_bullet(doc, "Run scenario tests: python -m unittest discover -s scenario_generator -p 'test_*.py'")
    add_bullet(doc, "Run ingestion tests: python -m unittest discover -s rag_pipeline -p 'test_*.py'")
    add_bullet(doc, "For every submitted call, confirm coherent pacing, sensible turn-taking, clear audio, active steering, transcript completeness, and a matching MP3.")
    add_bullet(doc, "Tie every reported bug to a transcript filename, timestamp, severity, actual behavior, expected behavior, and explanation of impact.")

    add_heading(doc, "8. Remaining submission actions", 1)
    remaining = [
        "Confirm at least ten output directories contain both a complete transcript and a playable MP3 or OGG recording.",
        "Select the strongest substantive failures and produce the final bug report with evidence references.",
        "Record and publish the project walkthrough (maximum three minutes) with webcam enabled.",
        "Record and publish the separate AI-assisted debugging walkthrough.",
        "Confirm repository visibility is public, run a secret scan, and verify .env is not committed.",
        "Submit the one originating phone number in E.164 format and verify it matches every assessment call.",
    ]
    for item in remaining:
        add_bullet(doc, item)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
